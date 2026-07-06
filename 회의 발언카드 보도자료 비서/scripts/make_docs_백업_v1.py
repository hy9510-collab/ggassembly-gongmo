# -*- coding: utf-8 -*-
"""
보도자료 / 발언 요약 보고서 → 한글 호환 문서(.docx) 생성
- .docx는 한글(아래아한글)에서 그대로 열리며, '다른 이름으로 저장 > hwp/hwpx'로 변환 가능합니다.

사용 예:
  python scripts/make_docs.py --press output/press/조미자_보도자료_data.json
  python scripts/make_docs.py --report output/press/조미자_보고서_data.json
  python scripts/make_docs.py --press p.json --report r.json
"""
import argparse
import json
import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from hwpx import HwpxDocument

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PRESS_DIR = os.path.join(BASE, "output", "press")

NAVY = RGBColor(0x18, 0x5F, 0xA5)
GRAY = RGBColor(0x66, 0x66, 0x66)
LGRAY = RGBColor(0x88, 0x88, 0x88)
CENTER = WD_ALIGN_PARAGRAPH.CENTER
RIGHT = WD_ALIGN_PARAGRAPH.RIGHT


def kfont(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def para(doc, text="", size=11, bold=False, align=None, color=None, after=8, line=1.5):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.line_spacing = line
    run = p.add_run(text)
    kfont(run, size=size, bold=bold, color=color)
    return p


def load(path):
    path = path if os.path.isabs(path) else os.path.join(BASE, path)
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def build_press(d):
    doc = Document()
    para(doc, d.get("org", "경기도의회"), size=14, bold=True, color=NAVY, after=2)
    para(doc, "보 도 자 료", size=10, color=GRAY, after=16)
    para(doc, d["title"], size=17, bold=True, align=CENTER, after=18, line=1.3)
    para(doc, d["lead"], size=11, after=12, line=1.7)
    for b in d.get("body", []):
        para(doc, b, size=11, after=10, line=1.7)
    para(doc, "", after=6)
    if d.get("date_line"):
        para(doc, d["date_line"], size=10, align=RIGHT, color=GRAY, after=2)
    if d.get("contact"):
        para(doc, d["contact"], size=9, color=LGRAY, after=0)
    return doc


def build_report(d):
    doc = Document()
    para(doc, d["title"], size=18, bold=True, align=CENTER, after=4)
    if d.get("subtitle"):
        para(doc, d["subtitle"], size=10.5, align=CENTER, color=GRAY, after=16)

    ov = d.get("overview", {})
    if ov:
        para(doc, "Ⅰ. 회의 개요", size=12.5, bold=True, color=NAVY, after=6)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for k, v in ov.items():
            cells = table.add_row().cells
            r0 = cells[0].paragraphs[0].add_run(k)
            kfont(r0, bold=True, size=10)
            r1 = cells[1].paragraphs[0].add_run(v)
            kfont(r1, size=10)
        para(doc, "", after=10)

    para(doc, "Ⅱ. 주요 발언 내용", size=12.5, bold=True, color=NAVY, after=8)
    for t in d.get("topics", []):
        para(doc, f"{t['no']}. {t['title']}", size=12, bold=True, after=4)
        if t.get("target"):
            para(doc, f"  · 대상기관: {t['target']}", size=10, color=GRAY, after=2, line=1.4)
        if t.get("summary"):
            para(doc, f"  · 발언요지: {t['summary']}", size=10.5, after=2, line=1.5)
        if t.get("quote"):
            para(doc, f"  · 핵심발언: “{t['quote']}”", size=10.5, color=NAVY, after=10, line=1.5)
        else:
            para(doc, "", after=6)

    if d.get("conclusion"):
        para(doc, "Ⅲ. 종합 의견", size=12.5, bold=True, color=NAVY, after=6)
        para(doc, d["conclusion"], size=11, after=8, line=1.7)

    if d.get("date_line"):
        para(doc, "", after=4)
        para(doc, d["date_line"], size=10, align=RIGHT, color=GRAY, after=0)
    return doc


# ── 한글 문서(.hwpx) 생성 (python-hwpx) ──
def _hp(doc, text, style, blank=False):
    doc.add_paragraph(str(text), char_pr_id_ref=style)
    if blank:
        doc.add_paragraph("")


def build_press_hwpx(d):
    doc = HwpxDocument.new()
    s_org = doc.ensure_run_style(bold=True, size=13, color="#185FA5")
    s_kind = doc.ensure_run_style(size=10, color="#666666")
    s_title = doc.ensure_run_style(bold=True, size=17, color="#16314A")
    s_body = doc.ensure_run_style(size=11)
    _hp(doc, d.get("org", "경기도의회"), s_org)
    _hp(doc, "보 도 자 료", s_kind, blank=True)
    _hp(doc, d["title"], s_title, blank=True)
    _hp(doc, d["lead"], s_body, blank=True)
    for b in d.get("body", []):
        _hp(doc, b, s_body, blank=True)
    if d.get("date_line"):
        _hp(doc, d["date_line"], s_kind)
    if d.get("contact"):
        _hp(doc, d["contact"], s_kind)
    return doc


def build_report_hwpx(d):
    doc = HwpxDocument.new()
    s_title = doc.ensure_run_style(bold=True, size=18, color="#16314A")
    s_sub = doc.ensure_run_style(size=10, color="#666666")
    s_head = doc.ensure_run_style(bold=True, size=13, color="#185FA5")
    s_topic = doc.ensure_run_style(bold=True, size=12)
    s_body = doc.ensure_run_style(size=11)
    s_quote = doc.ensure_run_style(size=10, color="#185FA5")
    _hp(doc, d["title"], s_title)
    if d.get("subtitle"):
        _hp(doc, d["subtitle"], s_sub)
    doc.add_paragraph("")
    ov = d.get("overview", {})
    if ov:
        _hp(doc, "Ⅰ. 회의 개요", s_head)
        for k, v in ov.items():
            _hp(doc, f"· {k}: {v}", s_body)
        doc.add_paragraph("")
    _hp(doc, "Ⅱ. 주요 발언 내용", s_head)
    for t in d.get("topics", []):
        _hp(doc, f"{t['no']}. {t['title']}", s_topic)
        if t.get("target"):
            _hp(doc, f"    · 대상기관: {t['target']}", s_body)
        if t.get("summary"):
            _hp(doc, f"    · 발언요지: {t['summary']}", s_body)
        if t.get("quote"):
            _hp(doc, f"    · 핵심발언: “{t['quote']}”", s_quote)
        doc.add_paragraph("")
    if d.get("conclusion"):
        _hp(doc, "Ⅲ. 종합 의견", s_head)
        _hp(doc, d["conclusion"], s_body)
    if d.get("date_line"):
        doc.add_paragraph("")
        _hp(doc, d["date_line"], s_sub)
    return doc


def main():
    ap = argparse.ArgumentParser(description="보도자료/요약보고서 DOCX 생성")
    ap.add_argument("--press", default=None, help="보도자료 데이터 JSON")
    ap.add_argument("--report", default=None, help="요약보고서 데이터 JSON")
    ap.add_argument("--outdir", default=PRESS_DIR)
    args = ap.parse_args()

    outdir = args.outdir if os.path.isabs(args.outdir) else os.path.join(BASE, args.outdir)
    os.makedirs(outdir, exist_ok=True)

    if args.press:
        d = load(args.press)
        base = f"보도자료_{d.get('member','의원')}_{d.get('tag','')}".rstrip("_")
        build_press(d).save(os.path.join(outdir, base + ".docx"))
        print("[완료] 보도자료(docx):", base + ".docx")
        build_press_hwpx(d).save_to_path(os.path.join(outdir, base + ".hwpx"))
        print("[완료] 보도자료(hwpx):", base + ".hwpx")

    if args.report:
        d = load(args.report)
        base = f"발언요약보고서_{d.get('member','의원')}_{d.get('tag','')}".rstrip("_")
        build_report(d).save(os.path.join(outdir, base + ".docx"))
        print("[완료] 요약보고서(docx):", base + ".docx")
        build_report_hwpx(d).save_to_path(os.path.join(outdir, base + ".hwpx"))
        print("[완료] 요약보고서(hwpx):", base + ".hwpx")

    if not args.press and not args.report:
        print("[안내] --press 또는 --report 로 데이터 JSON을 지정하세요.")


if __name__ == "__main__":
    main()
