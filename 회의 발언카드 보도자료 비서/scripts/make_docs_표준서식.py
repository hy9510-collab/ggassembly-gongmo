# -*- coding: utf-8 -*-
"""
보도자료 / 발언 요약 보고서 → docx + hwpx 생성
- 보도자료: 경기도의회 표준서식 (상단 표 + 제목 + ○부제 + 본문) 반영
- 요약보고서: 회의 개요표 + 주제별 발언 + 종합 의견
- .docx/.hwpx 모두 한글(아래아한글)에서 열립니다.

사용 예:
  python scripts/make_docs.py --press output/press/조미자_보도자료_data.json
  python scripts/make_docs.py --report output/press/조미자_보고서_data.json
"""
import argparse
import json
import os
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from hwpx import HwpxDocument

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PRESS_DIR = os.path.join(BASE, "output", "press")

NAVY = RGBColor(0x18, 0x5F, 0xA5)
GRAY = RGBColor(0x66, 0x66, 0x66)
CENTER = WD_ALIGN_PARAGRAPH.CENTER
RIGHT = WD_ALIGN_PARAGRAPH.RIGHT


def kfont(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rF = rPr.get_or_add_rFonts()
    rF.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def para(doc, text="", size=11, bold=False, align=None, color=None, after=8, line=1.5):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    r = p.add_run(text)
    kfont(r, size=size, bold=bold, color=color)
    return p


def set_cell(cell, text, bold=False, size=10, align=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(text)
    kfont(r, size=size, bold=bold, color=color)


def _hp(doc, text, style, blank=False):
    doc.add_paragraph(str(text), char_pr_id_ref=style)
    if blank:
        doc.add_paragraph("")


def load(path):
    path = path if os.path.isabs(path) else os.path.join(BASE, path)
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


# ──────────────────────────────────────────────
# 보도자료 (경기도의회 표준서식)
# ──────────────────────────────────────────────
def build_press(d):
    doc = Document()
    # 상단 표
    t = doc.add_table(rows=2, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    r0 = t.rows[0].cells
    set_cell(r0[0], "보 도 자 료", bold=True, size=15, align=CENTER, color=NAVY)
    set_cell(r0[1], "보 도 일 시", bold=True, size=10, align=CENTER)
    set_cell(r0[2], f"{d.get('release_date','')}   {d.get('distribute','배포 즉시')}", size=10, align=CENTER)
    r1 = t.rows[1].cells
    set_cell(r1[0], f"매수 {d.get('pages','')} · 참고자료 {d.get('ref_material','')} · 사진 {d.get('photo','')}", size=9, align=CENTER)
    set_cell(r1[1], "담 당", bold=True, size=10, align=CENTER)
    set_cell(r1[2], d.get("contact", ""), size=10, align=CENTER)

    para(doc, "", after=8)
    para(doc, d["title"], size=16, bold=True, align=CENTER, after=12, line=1.3)
    if d.get("subtitle"):
        para(doc, "○ " + d["subtitle"], size=12, bold=True, after=12)
    for b in d.get("body", []):
        para(doc, b, size=11, after=8, line=1.65)
    return doc


def build_press_hwpx(d):
    doc = HwpxDocument.new()
    s_title = doc.ensure_run_style(bold=True, size=16, color="#16314A")
    s_sub = doc.ensure_run_style(bold=True, size=12, color="#185FA5")
    s_body = doc.ensure_run_style(size=11)
    # 상단 표
    t = doc.add_table(2, 3)
    t.set_cell_text(0, 0, "보 도 자 료")
    t.set_cell_text(0, 1, "보 도 일 시")
    t.set_cell_text(0, 2, f"{d.get('release_date','')}   {d.get('distribute','배포 즉시')}")
    t.set_cell_text(1, 0, f"매수 {d.get('pages','')} · 참고자료 {d.get('ref_material','')} · 사진 {d.get('photo','')}")
    t.set_cell_text(1, 1, "담 당")
    t.set_cell_text(1, 2, d.get("contact", ""))
    doc.add_paragraph("")
    _hp(doc, d["title"], s_title)
    if d.get("subtitle"):
        _hp(doc, "○ " + d["subtitle"], s_sub)
    doc.add_paragraph("")
    for b in d.get("body", []):
        _hp(doc, b, s_body, blank=True)
    return doc


# ──────────────────────────────────────────────
# 발언 요약 보고서
# ──────────────────────────────────────────────
def build_report(d):
    doc = Document()
    para(doc, d["title"], size=18, bold=True, align=CENTER, after=4)
    if d.get("subtitle"):
        para(doc, d["subtitle"], size=10.5, align=CENTER, color=GRAY, after=16)
    ov = d.get("overview", {})
    if ov:
        para(doc, "Ⅰ. 회의 개요", size=12.5, bold=True, color=NAVY, after=6)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for k, v in ov.items():
            cells = table.add_row().cells
            r0 = cells[0].paragraphs[0].add_run(k)
            kfont(r0, bold=True, size=10)
            r1 = cells[1].paragraphs[0].add_run(v)
            kfont(r1, size=10)
        para(doc, "", after=10)
    para(doc, "Ⅱ. 주요 발언 내용", size=12.5, bold=True, color=NAVY, after=8)
    for t in d.get("topics", []):
        para(doc, f"{t['no']}. {t['title']}", size=12, bold=True, after=4)
        if t.get("target"):
            para(doc, f"  · 대상기관: {t['target']}", size=10, color=GRAY, after=2, line=1.4)
        if t.get("summary"):
            para(doc, f"  · 발언요지: {t['summary']}", size=10.5, after=2, line=1.5)
        if t.get("quote"):
            para(doc, f"  · 핵심발언: “{t['quote']}”", size=10.5, color=NAVY, after=10, line=1.5)
        else:
            para(doc, "", after=6)
    if d.get("conclusion"):
        para(doc, "Ⅲ. 종합 의견", size=12.5, bold=True, color=NAVY, after=6)
        para(doc, d["conclusion"], size=11, after=8, line=1.7)
    if d.get("date_line"):
        para(doc, "", after=4)
        para(doc, d["date_line"], size=10, align=RIGHT, color=GRAY, after=0)
    return doc


def build_report_hwpx(d):
    doc = HwpxDocument.new()
    s_title = doc.ensure_run_style(bold=True, size=18, color="#16314A")
    s_sub = doc.ensure_run_style(size=10, color="#666666")
    s_head = doc.ensure_run_style(bold=True, size=13, color="#185FA5")
    s_topic = doc.ensure_run_style(bold=True, size=12)
    s_body = doc.ensure_run_style(size=11)
    s_quote = doc.ensure_run_style(size=10, color="#185FA5")
    _hp(doc, d["title"], s_title)
    if d.get("subtitle"):
        _hp(doc, d["subtitle"], s_sub)
    doc.add_paragraph("")
    ov = d.get("overview", {})
    if ov:
        _hp(doc, "Ⅰ. 회의 개요", s_head)
        for k, v in ov.items():
            _hp(doc, f"· {k}: {v}", s_body)
        doc.add_paragraph("")
    _hp(doc, "Ⅱ. 주요 발언 내용", s_head)
    for t in d.get("topics", []):
        _hp(doc, f"{t['no']}. {t['title']}", s_topic)
        if t.get("target"):
            _hp(doc, f"    · 대상기관: {t['target']}", s_body)
        if t.get("summary"):
            _hp(doc, f"    · 발언요지: {t['summary']}", s_body)
        if t.get("quote"):
            _hp(doc, f"    · 핵심발언: “{t['quote']}”", s_quote)
        doc.add_paragraph("")
    if d.get("conclusion"):
        _hp(doc, "Ⅲ. 종합 의견", s_head)
        _hp(doc, d["conclusion"], s_body)
    if d.get("date_line"):
        doc.add_paragraph("")
        _hp(doc, d["date_line"], s_sub)
    return doc


def main():
    ap = argparse.ArgumentParser(description="보도자료/요약보고서 docx+hwpx 생성")
    ap.add_argument("--press", default=None)
    ap.add_argument("--report", default=None)
    ap.add_argument("--outdir", default=PRESS_DIR)
    args = ap.parse_args()

    outdir = args.outdir if os.path.isabs(args.outdir) else os.path.join(BASE, args.outdir)
    os.makedirs(outdir, exist_ok=True)

    if args.press:
        d = load(args.press)
        base = f"보도자료_{d.get('member','의원')}_{d.get('tag','')}".rstrip("_")
        build_press(d).save(os.path.join(outdir, base + ".docx"))
        print("[완료] 보도자료(docx):", base + ".docx")
        build_press_hwpx(d).save_to_path(os.path.join(outdir, base + ".hwpx"))
        print("[완료] 보도자료(hwpx):", base + ".hwpx")

    if args.report:
        d = load(args.report)
        base = f"발언요약보고서_{d.get('member','의원')}_{d.get('tag','')}".rstrip("_")
        build_report(d).save(os.path.join(outdir, base + ".docx"))
        print("[완료] 요약보고서(docx):", base + ".docx")
        build_report_hwpx(d).save_to_path(os.path.join(outdir, base + ".hwpx"))
        print("[완료] 요약보고서(hwpx):", base + ".hwpx")

    if not args.press and not args.report:
        print("[안내] --press 또는 --report 로 데이터 JSON을 지정하세요.")


if __name__ == "__main__":
    main()
