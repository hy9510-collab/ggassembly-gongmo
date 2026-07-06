# -*- coding: utf-8 -*-
"""
업로드된 회의록/발언 파일에서 텍스트 추출.
지원: .txt .docx .hwpx .pdf  (.hwp 구형 바이너리는 hwpx로 재저장 안내)

- 텍스트가 아주 길면 해당 의원 이름이 나오는 부분 중심으로 잘라
  AI 입력 크기를 적절히 유지한다.
"""
import os
import re
import sys
import zipfile

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

MAX_CHARS = 80000          # AI에 넘길 최대 글자 수
CONTEXT_LINES = 40         # 의원 이름 주변으로 남길 줄 수


def _read_txt(path):
    raw = open(path, "rb").read()
    for enc in ("utf-8-sig", "utf-8", "cp949"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _read_docx(path):
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def _read_hwpx(path):
    """hwpx = zip 안의 Contents/section*.xml — 태그 제거로 본문 추출."""
    texts = []
    with zipfile.ZipFile(path) as z:
        sections = sorted(n for n in z.namelist()
                          if re.match(r"Contents/section\d+\.xml$", n))
        for name in sections:
            xml = z.read(name).decode("utf-8", errors="replace")
            # 문단 경계를 줄바꿈으로
            xml = re.sub(r"</hp:p>", "\n", xml)
            body = re.sub(r"<[^>]+>", "", xml)
            body = (body.replace("&amp;", "&").replace("&lt;", "<")
                        .replace("&gt;", ">").replace("&quot;", '"'))
            texts.append(body)
    return "\n".join(texts)


def _read_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        text = _read_txt(path)
    elif ext == ".docx":
        text = _read_docx(path)
    elif ext == ".hwpx":
        text = _read_hwpx(path)
    elif ext == ".pdf":
        text = _read_pdf(path)
    elif ext == ".hwp":
        raise ValueError("구형 .hwp 파일은 읽을 수 없습니다. "
                         "한글에서 열어 [다른 이름으로 저장] → .hwpx 로 저장한 뒤 다시 올려주세요.")
    else:
        raise ValueError(f"지원하지 않는 파일 형식입니다: {ext} "
                         "(지원: txt, docx, hwpx, pdf)")
    text = re.sub(r"[ \t ]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text.strip()


def focus_member(text, member_name):
    """텍스트가 너무 길면 의원 이름 주변 블록만 남긴다."""
    if len(text) <= MAX_CHARS:
        return text, False
    lines = text.splitlines()
    keep = [False] * len(lines)
    hit = False
    for i, line in enumerate(lines):
        if member_name in line:
            hit = True
            for j in range(max(0, i - 5), min(len(lines), i + CONTEXT_LINES)):
                keep[j] = True
    if not hit:
        return text[:MAX_CHARS], True
    picked, prev = [], False
    for i, line in enumerate(lines):
        if keep[i]:
            picked.append(line)
            prev = True
        elif prev:
            picked.append("…(중략)…")
            prev = False
    out = "\n".join(picked)
    if len(out) > MAX_CHARS:
        out = out[:MAX_CHARS]
    return out, True


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("사용법: python read_transcript.py <파일> [의원이름]")
        sys.exit(1)
    t = extract_text(sys.argv[1])
    if len(sys.argv) > 2:
        t, trimmed = focus_member(t, sys.argv[2])
        print(f"[정보] 추출 {len(t):,}자 (잘림: {trimmed})")
    print(t[:2000])
